import psycopg2
import sys
import os
from os import system, name 
import csv
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import ns, parse_xml

from commodity import commodity
from measure import measure
from duty import duty
from quota_order_number import quota_order_number


class application(object):
	def __init__(self):
		self.clear()
		self.debug				= False
		self.country_codes		= ""

		# Initialise extended information
		self.geographical_area_name = ""
		self.agreement_title		= ""
		self.agreement_date			= ""
		self.version				= ""
		self.country_codes			= ""
		self.origin_quotas			= ""
		self.licensed_quota_volumes	= ""
		self.quota_scope			= ""
		self.quota_staging			= ""
		
		self.BASE_DIR			= os.path.dirname(os.path.abspath(__file__))
		self.SOURCE_DIR			= os.path.join(self.BASE_DIR, "source")
		self.CSV_DIR			= os.path.join(self.BASE_DIR, "csv")
		self.COMPONENT_DIR		= os.path.join(self.BASE_DIR, "xmlcomponents")
		self.CONFIG_DIR			= os.path.join(self.BASE_DIR, "..")
		self.CONFIG_DIR			= os.path.join(self.CONFIG_DIR, "..")
		self.CONFIG_DIR			= os.path.join(self.CONFIG_DIR, "create-data")
		self.CONFIG_DIR			= os.path.join(self.CONFIG_DIR, "config")
		self.CONFIG_FILE		= os.path.join(self.CONFIG_DIR, "config_common.json")
		self.CONFIG_FILE_LOCAL	= os.path.join(self.CONFIG_DIR, "config_migrate_measures_and_quotas.json")

		self.BALANCE_DIR		= os.path.join(self.BASE_DIR, "..")
		self.BALANCE_DIR		= os.path.join(self.BALANCE_DIR, "..")
		self.BALANCE_DIR		= os.path.join(self.BALANCE_DIR, "create-data")
		self.BALANCE_DIR		= os.path.join(self.BALANCE_DIR, "migrate_measures_and_quotas")
		self.BALANCE_DIR		= os.path.join(self.BALANCE_DIR, "source")
		self.BALANCE_DIR		= os.path.join(self.BALANCE_DIR, "quotas")
		self.BALANCE_FILE		= os.path.join(self.BALANCE_DIR, "quota_volume_master.csv")

		# For the word model folders
		self.MODEL_DIR			= os.path.join(self.BASE_DIR, "model")
		self.WORD_DIR			= os.path.join(self.MODEL_DIR, "word")
		self.DOCPROPS_DIR		= os.path.join(self.MODEL_DIR, "docProps")

		# For the output folders
		self.OUTPUT_DIR			= os.path.join(self.BASE_DIR, "output")

		self.get_config()


	def set_header(self, cell, part1, part2):
		grey = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(ns.nsdecls('w')))

		p = cell.paragraphs[0]
		p.add_run(part1).bold = True

		p = cell.add_paragraph()
		p.add_run(part2).bold = True

		cell.paragraphs[0].style = "Normal in Table"
		cell.paragraphs[1].style = "Normal in Table"
		#cell._tc.get_or_add_tcPr().grey


	def write_quota_section(self, document, section_text):
		table = document.add_table(rows = 1, cols = 8)
		table.style = "TableGrid"
		hdr_cells = table.rows[0].cells
		hdr = table.rows[0]

		self.set_header(hdr_cells[0], '(1)', 'Quota Number')
		self.set_header(hdr_cells[1], '(2)', 'Commodity Code')
		self.set_header(hdr_cells[2], '(3)', 'Country of Origin')
		self.set_header(hdr_cells[3], '(4)', 'Quota Duty Rate')
		self.set_header(hdr_cells[4], '(5)', 'Quota Volume')
		self.set_header(hdr_cells[5], '(6)', 'Date Quota Opens')
		self.set_header(hdr_cells[6], '(7)', 'Date Quota Closes')
		self.set_header(hdr_cells[7], '(8)', '2019 Quota Volume')

		cell_widths = [2.8, 2.9, 3, 4.5, 3, 3, 3, 3]
		for i in range(0, 7):
			hdr_cells[i].width = Cm(cell_widths[i])

		self.set_repeat_table_header(hdr)

		for q in self.quota_order_number_list:
			if q.section_text == section_text:
				my_row = table.add_row()
				row_cells = my_row.cells
				#self.set_cant_split(my_row)
				
				# Order number
				p = row_cells[0].paragraphs[0]
				cl = row_cells[0].paragraphs[0].add_run(q.ordernumber)
				p.style = "Normal in Table"

				# Commodity codes
				m = q.commodity_list[0]
				p = row_cells[1].paragraphs[0]
				p.add_run(m.commodity_code_formatted)
				p.style = "Normal in Table"
				if len(q.commodity_list) > 0:
					for i in range(1, len(q.commodity_list)):
						c = q.commodity_list[i]
						p = row_cells[1].add_paragraph(c.commodity_code_formatted)
						p.style = "Normal in Table"
				

				# Country of origin
				p = row_cells[2].paragraphs[0]
				p.add_run(q.format_country())
				p.style = "Normal in Table"

				# Quota duty rate
				# Need to check if these are all the same - if so, only print once
				unique_duty_rate = []
				for c in q.commodity_list:
					for m in c.measure_list:
						unique_duty_rate.append (m.duty_string)

				unique_duty_rate = set(unique_duty_rate)

				c = q.commodity_list[0]
				p = row_cells[3].paragraphs[0]
				if len(c.measure_list) > 0:
					p.add_run(c.measure_list[0].duty_string)
				else:
					p.add_run("")

				p.style = "Normal in Table"
				if len(unique_duty_rate) > 1:
					if len(q.commodity_list) > 0:
						for i in range(1, len(q.commodity_list)):
							c = q.commodity_list[i]
							if len(c.measure_list) > 0:
								p = row_cells[3].add_paragraph(c.combined_duty)
							else:
								p = row_cells[3].add_paragraph("")
							p.style = "Normal in Table"

				# quota volume
				p = row_cells[4].paragraphs[0]
				p.add_run("")
				p.style = "Normal in Table"

				# quota open date
				p = row_cells[5].paragraphs[0]
				p.add_run("")
				p.style = "Normal in Table"

				# quota close date
				p = row_cells[6].paragraphs[0]
				p.add_run("")
				p.style = "Normal in Table"

				# 2019 quota volume
				p = row_cells[7].paragraphs[0]
				p.add_run("")
				p.style = "Normal in Table"



	def set_cant_split(self, row):
		tr = row._tr
		trPr = tr.get_or_add_trPr()
		clause_cant_split = OxmlElement('w:cantSplit')
		clause_cant_split.set(qn('w:val'), "true")
		trPr.append(clause_cant_split)
		return row


	def set_repeat_table_header(self, row):
		tr = row._tr
		trPr = tr.get_or_add_trPr()
		tblHeader = OxmlElement('w:tblHeader')
		tblHeader.set(qn('w:val'), "true")
		trPr.append(tblHeader)
		return row


	def write_quotas(self):
		document = Document("template.docx")
		document.add_heading('The Customs (Tariff Quotas) (EU Exit) Regulations 2019 - United Kingdom Tariff Quotas, version 1.0, 22 March 2019', 0)

		document.add_heading('Part A - First Come First Served (FCFS)', 1)
		print ("Writing Part A")
		document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nunc eget congue purus. Vestibulum vestibulum tempor lectus sed facilisis. Nunc eu efficitur libero. Suspendisse a efficitur metus, et porta est. Integer pretium, est nec feugiat tempor, lacus nisl auctor nisi, sit amet congue ex nulla nec massa. Donec dapibus tempus nibh, ac lacinia risus euismod ut', style='Normal')
		self.write_quota_section(document, "Part A")

		document.add_heading('Part B: Section 1 - Authorised Use Provisions - FCFS', 1)
		print ("Writing Part B1")
		document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nunc eget congue purus. Vestibulum vestibulum tempor lectus sed facilisis. Nunc eu efficitur libero. Suspendisse a efficitur metus, et porta est. Integer pretium, est nec feugiat tempor, lacus nisl auctor nisi, sit amet congue ex nulla nec massa. Donec dapibus tempus nibh, ac lacinia risus euismod ut', style='Normal')
		self.write_quota_section(document, "Part B - Section 1")

		document.add_heading('Part B: Section 2 - Authorised Use Provisions Fisheries', 1)
		print ("Writing Part B2")
		document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nunc eget congue purus. Vestibulum vestibulum tempor lectus sed facilisis. Nunc eu efficitur libero. Suspendisse a efficitur metus, et porta est. Integer pretium, est nec feugiat tempor, lacus nisl auctor nisi, sit amet congue ex nulla nec massa. Donec dapibus tempus nibh, ac lacinia risus euismod ut', style='Normal')
		self.write_quota_section(document, "Part B - Section 2")

		document.add_heading('Part C: Section 1 - Licence Managed Quotas', 1)
		print ("Writing Part C1")
		document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nunc eget congue purus. Vestibulum vestibulum tempor lectus sed facilisis. Nunc eu efficitur libero. Suspendisse a efficitur metus, et porta est. Integer pretium, est nec feugiat tempor, lacus nisl auctor nisi, sit amet congue ex nulla nec massa. Donec dapibus tempus nibh, ac lacinia risus euismod ut', style='Normal')
		self.write_quota_section(document, "Part C - Section 1")

		document.add_heading('Part C: Section 2 - Authorised Use Provisions - Licence Managed Quotas', 1)
		print ("Writing Part C2")
		document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nunc eget congue purus. Vestibulum vestibulum tempor lectus sed facilisis. Nunc eu efficitur libero. Suspendisse a efficitur metus, et porta est. Integer pretium, est nec feugiat tempor, lacus nisl auctor nisi, sit amet congue ex nulla nec massa. Donec dapibus tempus nibh, ac lacinia risus euismod ut', style='Normal')
		self.write_quota_section(document, "Part C - Section 2")


		

		document.save('demo.docx')


	def get_quotas(self):

		quota_list = "'090006', '090007', '090008', '090009', '090013', '090023', '090027', '090035', '090039', '090040', '090041', '090043', '090046', '090047', '090053', '090055', '090056', '090058', '090059', '090060', '090061', '090062', '090063', '090067', '090070', '090071', '090072', '090073', '090076', '090083', '090085', '090086', '090088', '090089', '090092', '090093', '090094', '090095', '090096', '090097', '090098', '090101', '090103', '090104', '090106', '090107', '090109', '090111', '090118', '090120', '090123', '090124', '090126', '090127', '090128', '090130', '090131', '090132', '090693', '090706', '090707', '090708', '090790', '091922', '092003', '092011', '092012', '092013', '092014', '092016 ', '092017', '092201', '092501', '092600', '092602', '092617', '092620', '092624', '092628', '092629', '092631', '092633', '092634', '092637', '092638', '092639', '092641', '092643', '092644', '092645', '092646', '092647', '092648', '092649', '092650', '092652', '092658', '092659', '092661', '092662', '092664', '092665', '092668', '092671', '092672', '092673', '092674', '092675', '092676', '092679', '092681', '092682', '092683', '092684', '092685', '092686', '092688', '092690', '092694', '092696', '092697', '092698', '092700', '092704', '092708', '092710', '092720', '092721', '092722', '092723', '092728', '092730', '092734', '092736', '092738', '092740', '092742', '092746', '092748', '092750', '092754', '092759', '092760', '092761', '092762', '092763', '092765', '092769', '092770', '092772', '092774', '092776', '092777', '092778', '092784', '092785', '092786', '092788', '092790', '092794', '092798', '092799', '092800', '092802', '092804', '092806', '092808', '092812', '092814', '092816', '092820', '092822', '092824', '092826', '092828', '092829', '092830', '092832', '092834', '092835', '092837', '092840', '092842', '092846', '092848', '092849', '092850', '092851', '092852', '092854', '092856', '092858', '092860', '092864', '092866', '092868', '092870', '092872', '092874', '092876', '092878', '092880', '092888', '092889', '092906', '092907', '092908', '092909', '092910', '092913', '092928', '092929', '092932', '092933', '092935', '092945', '092955', '092972', '092975', '094001', '094002', '094003', '094015', '094038', '094057', '094058', '094067', '094068', '094069', '094105', '094106', '094112', '094113', '094116', '094117', '094118', '094119', '094123', '094124', '094125', '094126', '094127', '094128', '094129', '094130', '094131', '094148', '094150', '094153', '094154', '094166', '094168', '094170', '094195', '094204', '094211', '094212', '094213', '094214', '094215', '094216 ', '094217', '094218', '094251', '094252', '094253', '094254', '094255', '094256', '094257', '094258', '094259', '094260', '094261', '094262', '094263', '094264', '094265', '094317', '094318', '094320', '094321', '094410', '094411', '094412', '094420', '094421', '094422', '094450', '094451', '094452', '094453', '094454', '094513', '094514', '094515', '094590', '094594', '094595', '094700', '094701', '097701', '097702', '097703', '097704', '097705', '097706', '097707', '097708', '097709', '097710', '097711', '097712', '097713', '098700', '098701', '098702', '098703', '098704', '098705', '098707', '098708', '098709', '908706'"


		# Get the unique commodities
		sql = """
		select distinct section_text, ordernumber, m.goods_nomenclature_item_id, m.geographical_area_id,
		g.description as country
		from ml.v5_2019 m, ml.ml_geographical_areas g, ml.quota_section q,
		goods_nomenclatures gn
		where gn.goods_nomenclature_item_id = m.goods_nomenclature_item_id
		and gn.validity_end_date is null
		and m.geographical_area_id = g.geographical_area_id
		and q.quota_order_number_id = m.ordernumber
		and m.measure_type_id in ('122', '123')
		and ordernumber in (""" + quota_list + """
		)
		order by 1, 2, 3
		"""

		cur = self.conn_eu.cursor()
		cur.execute(sql)
		rows = cur.fetchall()
		self.commodity_list = []
		for rw in rows:
			section_text					= rw[0]
			ordernumber						= rw[1]
			goods_nomenclature_item_id		= rw[2]
			geographical_area_id			= rw[3]
			country							= rw[4]
			m = commodity(section_text, ordernumber, goods_nomenclature_item_id, geographical_area_id, country)

			self.commodity_list.append(m)


		# Get the unique measures
		sql = """
		select distinct section_text, ordernumber, m.goods_nomenclature_item_id, mc.*, m.validity_start_date, m.validity_end_date
		from ml.v5_2019 m, measure_components mc, ml.quota_section q, goods_nomenclatures gn
		where m.measure_sid = mc.measure_sid
		and gn.goods_nomenclature_item_id = m.goods_nomenclature_item_id
		and gn.validity_end_date is null
		and q.quota_order_number_id = m.ordernumber
		and m.measure_type_id in ('122', '123')
		and ordernumber in (""" + quota_list + """)
		and duty_expression_id in ('01', '04', '19', '20', '15', '17', '35')
		order by 1, 2, m.validity_start_date DESC, mc.duty_expression_id
		"""

		cur = self.conn_eu.cursor()
		cur.execute(sql)
		rows = cur.fetchall()
		self.measure_list = []
		for rw in rows:
			section_text					= rw[0]
			ordernumber						= rw[1]
			goods_nomenclature_item_id		= rw[2]
			measure_sid						= rw[3]
			duty_expression_id				= rw[4]
			duty_amount						= rw[5]
			monetary_unit_code				= rw[6]
			measurement_unit_code			= rw[7]
			measurement_unit_qualifier_code	= rw[8]
			validity_start_date				= rw[9]
			validity_end_date				= rw[10]

			m = measure(section_text, ordernumber, goods_nomenclature_item_id, measure_sid, duty_expression_id, duty_amount, monetary_unit_code, measurement_unit_code, measurement_unit_qualifier_code, validity_start_date, validity_end_date)
			self.measure_list.append(m)

		# This creates a full list of all of the measures that meet the criteria that have an end date after January 2018
		# We want to remove duplicates, which there surely will be
		concat_list = []
		for m in self.measure_list:
			if m.concatentated_fields in concat_list:
				m.mark_for_deletion = True
				#print ("Marking for deletion")
			else:
				concat_list.append (m.concatentated_fields)

		for i in range(len(self.measure_list)-1, -1, -1):
			m = self.measure_list[i]
			if m.mark_for_deletion == True:
				self.measure_list.pop(i)
				#print ("Popping")

		#sys.exit()
		quota_order_number_list  = []
		self.quota_order_number_list = []
		for m in self.commodity_list:
			quota_order_number_list.append (m.ordernumber)
		
		quota_order_number_list = set(quota_order_number_list)
		quota_order_number_list = sorted(quota_order_number_list)

		for q in quota_order_number_list:
			q2 = quota_order_number(q)
			self.quota_order_number_list.append(q2)

		# Append measures to commodities		
		for c in self.commodity_list:
			for m in self.measure_list:
				if m.goods_nomenclature_item_id == c.goods_nomenclature_item_id:
					a = 1
					if m.ordernumber == c.ordernumber:
						c.measure_list.append(m)
						#break

		for c in self.commodity_list:
			c.combine_duty_strings()
			"""
			if len(c.measure_list) > 1:
				print ("More than one", c.goods_nomenclature_item_id, str(len(c.measure_list)), c.measure_list[0].duty_string, c.measure_list[1].duty_string)
				sys.exit()
			"""


		# Append commodities to quota order numbers
		for c in self.commodity_list:
			for q in self.quota_order_number_list:
				if c.ordernumber == q.ordernumber:
					q.section_text = c.section_text
					q.country = c.country
					q.commodity_list.append(c)
					break



	def clear(self): 
		# for windows 
		if name == 'nt': 
			_ = system('cls') 
		# for mac and linux(here, os.name is 'posix') 
		else: 
			_ = system('clear')


	def get_config(self):
		# Get global config items
		with open(self.CONFIG_FILE, 'r') as f:
			my_dict = json.load(f)

		#self.DBASE	= my_dict['dbase']
		#self.DBASE	= "tariff_eu"
		
		self.DBASE_EU	= my_dict['dbase_eu']
		self.DBASE_UK	= my_dict['dbase_uk']
		
		self.p		= my_dict['p']

		# Get local config items
		with open(self.CONFIG_FILE_LOCAL, 'r') as f:
			my_dict = json.load(f)

		self.all_country_profiles = my_dict['country_profiles']

		# Connect to the database
		self.connect()


	def connect(self):
		self.conn_uk = psycopg2.connect("dbname=" + self.DBASE_UK + " user=postgres password=" + self.p)
		self.conn_eu = psycopg2.connect("dbname=" + self.DBASE_EU + " user=postgres password=" + self.p)


	def shutDown(self):
		self.conn_uk.close()
		self.conn_eu.close()
